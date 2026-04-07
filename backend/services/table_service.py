from docx import Document
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
 
 
def replace_ctc_table_dynamic(doc, df, section="initial"):
    if df is None or df.empty:
        print(f"{section} sheet is empty")
        return
 
    print(f"📊 Processing {section.upper()} | Tables:", len(doc.tables))
 
    # ==============================
    # HELPERS
    # ==============================
    def clean_text(text):
        text = str(text).lower()
        text = re.sub(r'[^a-z0-9\s]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text
 
    def clean_value(val):
        if str(val).lower() == "nan" or val == "":
            return ""
        try:
            return str(int(float(str(val).replace(",", ""))))
        except:
            return str(val).strip()
 
    def apply_highlight(cell):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '808080')  # grey
        cell._tc.get_or_add_tcPr().append(shading)
 
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
 
    # ==============================
    # SPLIT DATA (NO MIXING)
    # ==============================
    ctc_data = []
    retirals_data = []
 
    is_retiral_section = False
 
    for _, row in df.iterrows():
        values = [clean_value(v) for v in row]
 
        # Skip empty rows
        if all(v == "" for v in values):
            continue
 
        # Skip rows with no label
        if values[0] == "":
            continue
 
        row_text = clean_text(" ".join(values))
 
        # Skip unwanted rows
        if "annual ctc details" in row_text:
            continue
 
        if "esic" in row_text:
            continue
 
        # Section switch
        if "retirals" in row_text:
            is_retiral_section = True
            continue
 
        if not is_retiral_section:
            ctc_data.append(values)
        else:
            retirals_data.append(values)
 
    print("CTC ROWS:", len(ctc_data))
    print("RETIRALS ROWS:", len(retirals_data))
 
    # ==============================
    # TABLE SELECTION
    # ==============================
    try:
        if section == "initial":
            ctc_table = doc.tables[0]
            retirals_table = doc.tables[1]
        else:
            if len(doc.tables) < 4:
                print("Revised tables missing in template")
                return
            ctc_table = doc.tables[2]
            retirals_table = doc.tables[3]
    except Exception as e:
        print("Table selection error:", e)
        return
 
    # ==============================
    # REMOVE HEADER FROM RETIRALS (WORD)
    # ==============================
    if retirals_table and len(retirals_table.rows) > 0:
        retirals_table._element.remove(retirals_table.rows[0]._element)
 
    # ==============================
    # FILL TABLE FUNCTION
    # ==============================
    def fill_table(table, data, is_retiral=False):
        if not table:
            print(f"Table not found for {section}")
            return
 
        if not data:
            print(f"No data for {section}")
            return
 
        # Keep only header row (for CTC)
        if is_retiral:
            #  Remove ALL rows (including header)
            while len(table.rows) > 0:
                table._element.remove(table.rows[0]._element)
        else:
            #  Keep header for CTC table
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
 
        for idx, row in enumerate(data):
            new_cells = table.add_row().cells
 
            for i in range(len(new_cells)):
                value = row[i] if i < len(row) else ""
                new_cells[i].text = value
 
            # Bold first column
            if len(new_cells) > 0:
                for paragraph in new_cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
 
            row_text = clean_text(" ".join(row))
 
            # Highlight CTC total
            if not is_retiral:
                if "annual gross salary" in row_text:
                    for cell in new_cells:
                        apply_highlight(cell)
 
            # Highlight last retiral row
            if is_retiral:
                if idx == len(data) - 1:
                    for cell in new_cells:
                        apply_highlight(cell)
 
    # ==============================
    # APPLY TABLES
    # ==============================
    fill_table(ctc_table, ctc_data, is_retiral=False)
    fill_table(retirals_table, retirals_data, is_retiral=True)
 
    print(f"✅ {section.upper()} tables filled successfully!")"